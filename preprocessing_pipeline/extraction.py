from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
import subprocess
import tempfile
from typing import Iterable, List

from docx import Document
import olefile
from striprtf.striprtf import rtf_to_text


SUPPORTED_EXTENSIONS = {".docx", ".doc", ".rtf"}
_NOISE_MARKERS = (
    "thememanager.xml",
    "word.document.8",
    "fonttbl",
    "objdata",
    "xml version",
    "times new roman",
    "arial",
    "calibri",
    "aptos",
    "symbol",
)


@dataclass(frozen=True)
class SourceFile:
    path: Path
    year: int
    reviewer_id: str


def should_skip_file(path: Path) -> bool:
    if path.name.startswith("~$"):
        return True
    if path.name == ".DS_Store":
        return True
    return path.suffix.lower() not in SUPPORTED_EXTENSIONS


def parse_metadata_from_path(path: Path) -> tuple[int, str]:
    # Expected: .../Отзывы_YYYY/YYYY_RN/filename.ext
    round_dir = path.parent.name
    year_dir = path.parent.parent.name if path.parent.parent else ""

    year = None
    if year_dir.startswith("Отзывы_"):
        year_str = year_dir.replace("Отзывы_", "", 1)
        if year_str.isdigit():
            year = int(year_str)
    if year is None:
        prefix = round_dir.split("_", 1)[0]
        if prefix.isdigit():
            year = int(prefix)
    if year is None:
        raise ValueError(f"Cannot parse year from path: {path}")

    reviewer_id = ""
    if "_R" in round_dir:
        reviewer_id = round_dir.split("_R", 1)[1]
    if not reviewer_id:
        raise ValueError(f"Cannot parse reviewer_id from path: {path}")

    return year, f"R{reviewer_id}"


def discover_source_files(input_dir: Path) -> List[SourceFile]:
    all_files = [p for p in input_dir.rglob("*") if p.is_file()]
    filtered = [p for p in all_files if not should_skip_file(p)]

    deduped = deduplicate_by_stem(filtered)
    result: List[SourceFile] = []
    for path in deduped:
        year, reviewer_id = parse_metadata_from_path(path)
        result.append(SourceFile(path=path, year=year, reviewer_id=reviewer_id))
    return sorted(result, key=lambda item: str(item.path))


def deduplicate_by_stem(paths: Iterable[Path]) -> List[Path]:
    # Keep best format per stem in the same directory: .docx > .doc > .rtf
    priority = {".docx": 3, ".doc": 2, ".rtf": 1}
    selected: dict[tuple[str, str], Path] = {}
    for path in paths:
        key = (str(path.parent), path.stem.lower())
        current = selected.get(key)
        if current is None:
            selected[key] = path
            continue
        if priority.get(path.suffix.lower(), 0) > priority.get(current.suffix.lower(), 0):
            selected[key] = path
    return list(selected.values())


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".rtf":
        return extract_rtf(path)
    if suffix == ".doc":
        return extract_doc_legacy(path)
    raise ValueError(f"Unsupported file extension: {path}")


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    table_chunks: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                table_chunks.append(" | ".join(cells))
    chunks = paragraphs + table_chunks
    return "\n".join(chunks).strip()


def extract_rtf(path: Path) -> str:
    raw = path.read_text(encoding="latin-1", errors="ignore")
    try:
        text = rtf_to_text(raw)
    except Exception as exc:  # pragma: no cover - defensive path
        raise RuntimeError(f"Cannot parse RTF file: {path}") from exc
    cleaned = _normalize_extracted_text(text)
    cleaned = _strip_binary_noise(cleaned)
    if _looks_like_binary_garbage(cleaned):
        soffice_text = _extract_with_soffice_to_text(path)
        if soffice_text:
            cleaned = soffice_text
    if _looks_like_binary_garbage(cleaned):
        raise RuntimeError(f"Extracted .rtf text looks like binary garbage: {path}")
    return cleaned


def extract_doc_legacy(path: Path) -> str:
    soffice_text = _extract_with_soffice_to_text(path)
    if soffice_text:
        return soffice_text

    data = path.read_bytes()
    if data.startswith(b"PK\x03\x04"):
        # Some sources are .docx renamed to .doc.
        return extract_docx(path)

    if not olefile.isOleFile(str(path)):
        raise RuntimeError(f"Legacy .doc is not a valid OLE2 file: {path}")

    stream_data = _read_doc_ole_stream_data(path)
    candidates: list[str] = []
    for decoded in _decode_doc_stream_candidates(stream_data):
        chunks = _extract_text_chunks(decoded)
        if not chunks:
            continue
        candidate = _normalize_extracted_text("\n".join(chunks))
        candidate = _strip_binary_noise(candidate)
        if candidate and not _looks_like_binary_garbage(candidate):
            candidates.append(candidate)

    if not candidates:
        raise RuntimeError(f"Cannot extract readable text from legacy .doc: {path}")
    return max(candidates, key=_score_extracted_text)


def _extract_with_soffice_to_text(path: Path) -> str:
    with tempfile.TemporaryDirectory(prefix="textmining_soffice_") as temp_dir:
        try:
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    temp_dir,
                    str(path),
                ],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                timeout=180,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, FileNotFoundError):
            return ""

        txt_candidates = sorted(Path(temp_dir).glob("*.txt"))
        if not txt_candidates:
            return ""
        text = txt_candidates[0].read_text(encoding="utf-8", errors="ignore")
        cleaned = _normalize_extracted_text(text)
        cleaned = _strip_binary_noise(cleaned)
        return "" if _looks_like_binary_garbage(cleaned) else cleaned


def _read_doc_ole_stream_data(path: Path) -> bytes:
    parts: list[bytes] = []
    with olefile.OleFileIO(str(path)) as ole:
        for stream_name in (["WordDocument"], ["1Table"], ["0Table"]):
            if not ole.exists(stream_name):
                continue
            parts.append(ole.openstream(stream_name).read())
    if not parts:
        raise RuntimeError(f"Legacy .doc has no readable Word streams: {path}")
    return b"\n".join(parts)


def _decode_doc_stream_candidates(stream_data: bytes) -> list[str]:
    candidates: list[str] = []
    utf16_slice = stream_data[: len(stream_data) - (len(stream_data) % 2)]
    for enc, payload in (
        ("utf-16le", utf16_slice),
        ("cp1251", stream_data),
        ("utf-8", stream_data),
        ("latin-1", stream_data),
    ):
        decoded = payload.decode(enc, errors="ignore")
        if decoded:
            candidates.append(decoded)
    return candidates


def _extract_text_chunks(text: str) -> list[str]:
    cleaned = re.sub(r"[\x00-\x08\x0B-\x1F\x7F]", " ", text.replace("\r", "\n"))
    pieces = re.split(r"[\n\t]+", cleaned)
    chunks: list[str] = []
    for piece in pieces:
        piece = re.sub(r"\s+", " ", piece).strip()
        if len(piece) < 20:
            continue
        lowered = piece.lower()
        if any(
            marker in lowered
            for marker in (
                "fonttbl",
                "colortbl",
                "stylesheet",
                "worddocument",
                "thememanager.xml",
                "objdata",
                "officedocument",
                "xml version",
            )
        ):
            continue
        letters = sum(ch.isalpha() for ch in piece)
        if letters / max(1, len(piece)) < 0.45:
            continue
        if not re.search(r"[А-Яа-яЁёA-Za-z]{3,}", piece):
            continue
        chunks.append(piece)

    # Preserve order while deduplicating near-identical chunks.
    seen: set[str] = set()
    unique: list[str] = []
    for chunk in chunks:
        key = chunk.lower()
        if key in seen:
            continue
        seen.add(key)
        unique.append(chunk)
    return unique


def _normalize_extracted_text(text: str) -> str:
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines).strip()


def _strip_binary_noise(text: str) -> str:
    # Keep only chars typical for natural-language reviews; drop font/binary leftovers.
    kept_lines: list[str] = []
    for raw_line in text.splitlines():
        line = _sanitize_line(raw_line)
        if not line:
            continue
        lowered = line.lower()
        if any(marker in lowered for marker in _NOISE_MARKERS):
            continue
        kept_lines.append(line)
    return "\n".join(kept_lines).strip()


def _sanitize_line(line: str) -> str:
    # Replace non-text Unicode noise while preserving punctuation and digits.
    line = re.sub(r"[^A-Za-zА-Яа-яЁё0-9\s\.,;:!?\-–—()\"'«»/№%+=&…]", " ", line)
    # Remove long runs of isolated one-char tokens typical for binary debris.
    line = re.sub(r"(?:\b[0-9A-Za-z]\b(?:\s+|$)){6,}", " ", line)
    line = re.sub(r"\s+", " ", line).strip()
    if len(line) < 3:
        return ""
    letters = len(re.findall(r"[A-Za-zА-Яа-яЁё]", line))
    words = len(re.findall(r"[A-Za-zА-Яа-яЁё]{3,}", line))
    if letters == 0 and not re.search(r"\d", line):
        return ""
    if words == 0 and letters < 6:
        return ""
    # Drop mostly symbolic lines that survived replacement.
    if letters / max(1, len(line)) < 0.28 and not re.search(r"[А-Яа-яЁё]{3,}", line):
        return ""
    return line


def _score_extracted_text(text: str) -> tuple[int, int, int]:
    cyr = len(re.findall(r"[А-Яа-яЁё]", text))
    letters = len(re.findall(r"[A-Za-zА-Яа-яЁё]", text))
    words = len(re.findall(r"[A-Za-zА-Яа-яЁё]{3,}", text))
    return (cyr, words, letters)


def _looks_like_binary_garbage(text: str) -> bool:
    if not text:
        return True
    lowered = text.lower()
    if (
        "thememanager.xml" in lowered
        or "word.document.8" in lowered
        or "fonttbl" in lowered
        or "objdata" in lowered
        or "xml version" in lowered
    ):
        return True
    if text.count("яяя") > 20:
        return True
    letters = sum(ch.isalpha() for ch in text)
    if len(text) >= 400 and letters / max(1, len(text)) < 0.30:
        return True
    return False

